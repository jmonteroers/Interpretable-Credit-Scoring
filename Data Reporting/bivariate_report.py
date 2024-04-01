"""
This script takes the train data after feature engineering (train_apps_ext) and creates two types of editable reports:

- First report where we count the number of missing observations for each feature- If there are any, then create a barplot that compares the % of bad for missing vs non-missing
- Second report where, for categorical variables a barplot is created with % bads y-axis, x-axis categories (incl. missing); for quantitative variables, a density plot broken down by good/bads (following Carrasco, 2023) - excluding missing
"""
import docx
from io import BytesIO
import pandas as pd
import matplotlib.pyplot as plt

TARGET = "TARGET"


def add_section_missing(doc, df, feat_name):
    # Get number of missing
    n_missing = df[feat_name].isna().sum()

    # Create barplot (bytes)
    barplot = create_barplot_missing(df, feat_name) if n_missing > 0 else None

    # Add to Doc file
    doc.add_heading(feat_name)
    doc.add_paragraph(f"The number of missing observations for {feat_name} is: {n_missing}")
    if barplot is not None:
        doc.add_picture(barplot, width=docx.shared.Inches(5))


def create_barplot_missing(df, feature):
    """Returns barplot as a BytesIO Object"""
    # create copy, to prevent changes to original df
    df = df.copy()
    # create feature with missing/non-missing
    na_feat = f"{feature}_na"
    df[na_feat] = "Non-missing"
    df.loc[df[feature].isna(), na_feat] = "Missing"
    # aggregate by feature missing/non-missing
    agg_df = df.groupby(na_feat)[[TARGET]].mean().reset_index()
    agg_df.rename(columns={TARGET: "% Bad"}, inplace=True)

    # create barplot
    barplot_image = BytesIO()
    ax = agg_df.plot.bar(x = na_feat, y = "% Bad", color = "#43ff64d9")
    ax.set_xlabel(feature.capitalize())
    ax.set_ylabel("Proportion Bad")
    # horizontal x-axis labels
    plt.xticks(rotation=0)
    plt.savefig(barplot_image, format="png")
    plt.close()
    # reset pointer bytes to starting point
    barplot_image.seek(0)

    return barplot_image


def create_report_missing(df, output_path):
    # Create Word
    doc = docx.Document()

    # Define List of Variables to Include
    # Sorted in decreasing number of missing variables
    na_counts = df.isna().sum()
    sorted_columns = na_counts.sort_values(ascending=False)
    sorted_column_names = sorted_columns.index.tolist()
    excl_vars = ["SK_ID_CURR", "TARGET"]
    report_vars = [col for col in sorted_column_names if col not in excl_vars]

    # add to Word file
    for var in report_vars:
        add_section_missing(doc, df, var)
    
    # Save Word
    doc.save(output_path)

if __name__ == "__main__":
    from pathlib import Path
    PARENT_DIR = Path(__file__).absolute().parents[2] / 'Data' / 'Home Credit'
    df = pd.read_csv(PARENT_DIR / 'processed' / 'train_apps_ext.csv.zip')

    create_report_missing(df, PARENT_DIR / "meta" / "bivariate_report_missing.docx")



